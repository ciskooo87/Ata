import streamlit as st
import whisper
from docx import Document
import tempfile
import os

# Carrega modelo Whisper (pequeno para rodar mais rápido)
@st.cache_resource

def load_model():
    return whisper.load_model("small")  # troque para "medium" ou "large" se quiser mais precisão

model = load_model()

st.title("📄 Gerador de Atas a partir de Áudio")

uploaded_file = st.file_uploader("Envie o arquivo de áudio (MP3/WAV/M4A)", type=["mp3", "wav", "m4a"])

if uploaded_file:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp:
        tmp.write(uploaded_file.read())
        tmp_path = tmp.name

    st.info("🔄 Transcrevendo o áudio, pode levar alguns minutos...")
    result = model.transcribe(tmp_path, fp16=False, language="pt")  # força português
    transcription = result["text"]

    st.subheader("📝 Transcrição Completa")
    st.write(transcription)

    st.subheader("📌 Resumo Automático")
    resumo = transcription[:1000] + ("..." if len(transcription) > 1000 else "")
    st.write(resumo)

    # Coleta dados obrigatórios
    st.subheader("📋 Informações da Ata")
    data = st.date_input("Data da Reunião")
    participantes = st.text_area("Participantes (separados por vírgula)")
    pauta = st.text_area("Pauta da Reunião")

    if st.button("Gerar ATA em DOCX"):
        doc = Document()
        doc.add_heading("ATA DE REUNIÃO", level=1)

        doc.add_paragraph(f"Data: {data}")
        doc.add_paragraph(f"Participantes: {participantes}")
        doc.add_paragraph(f"Pauta: {pauta}")

        doc.add_heading("Transcrição", level=2)
        doc.add_paragraph(transcription)

        doc.add_heading("Resumo", level=2)
        doc.add_paragraph(resumo)

        output_path = os.path.join(tempfile.gettempdir(), "ATA.docx")
        doc.save(output_path)

        with open(output_path, "rb") as f:
            st.download_button("📥 Baixar ATA em DOCX", f, file_name="ATA.docx")
